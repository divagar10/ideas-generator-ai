import google.generativeai as genai
import docx
from datetime import datetime

# Set your API key
genai.configure(api_key="AIzaSyDcc1s1OF-U7Dr6n4s5MVzHJbzyzfZxVrY")

# Use a capable model
model = genai.GenerativeModel('models/gemini-1.5-flash')

# Inputs
platform = input("Enter platform (YouTube / Instagram / Blog): ")
niche = input("Enter your content niche: ")
audience_type = input("Who is your target audience? (e.g., youth, professionals): ")

# Prompt Engineering
prompt = f"""
You are a creative AI assistant for content creators. Generate 10 unique and viral content ideas for {platform} creators in the "{niche}" niche. 
Each idea should be engaging, suitable for the {audience_type}, and relevant to current trends.
Output the ideas as a numbered list.
"""

# Generate Response
response = model.generate_content(prompt)
ideas = response.text.strip()
print("\nGenerated Content Ideas:\n")
print(ideas)

# Export to .txt
filename_txt = f"content_ideas_{platform}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
with open(filename_txt, "w", encoding="utf-8") as file:
    file.write(ideas)
print(f"\n✅ Exported to {filename_txt}")

# Export to .docx
doc = docx.Document()
doc.add_heading(f'Content Ideas for {platform} - {niche}', 0)
doc.add_paragraph(ideas)
filename_docx = filename_txt.replace('.txt', '.docx')
doc.save(filename_docx)
print(f"✅ Exported to {filename_docx}")
